import hashlib
import json
import math
import re
import time
from collections import Counter
from pathlib import Path
from threading import Lock

from backend.core import workspace

ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
INDEX_FILE = DATA_DIR / "rag_index.json"

SKIP_DIRS = {".git", "__pycache__", "node_modules", ".venv", "venv"}
SKIP_EXTS = {".pyc", ".log", ".png", ".jpg", ".jpeg", ".gif", ".webp", ".ico", ".exe", ".dll", ".zip"}
TEXT_EXTS = {
    ".txt", ".md", ".py", ".js", ".html", ".css", ".json", ".csv", ".tsv",
    ".yml", ".yaml", ".toml", ".ini", ".sql", ".xml", ".svg"
}
DOCUMENT_EXTS = {".pdf", ".docx", ".xlsx"}
VECTOR_DIMS = 384
MAX_TEXT_BYTES = 1_500_000
MAX_DOCUMENT_BYTES = 12_000_000
MAX_XLSX_ROWS_PER_SHEET = 5000

_lock = Lock()


def _blank():
    return {"version": 2, "updated_at": None, "docs": [], "chunks": []}


from backend.core import runtime_store as store

def _load():
    data = store.get_kv("rag_vector")
    if not isinstance(data, dict):
        DATA_DIR.mkdir(exist_ok=True)
        if INDEX_FILE.exists():
            with _lock:
                try:
                    data = json.loads(INDEX_FILE.read_text(encoding="utf-8"))
                except Exception:
                    data = _blank()
        else:
            data = _blank()
        store.set_kv("rag_vector", data)
    if data.get("version") != 2:
        return _blank()
    return data


def raw_index():
    return _load()


def _save(index):
    DATA_DIR.mkdir(exist_ok=True)
    index["updated_at"] = time.time()
    with _lock:
        store.set_kv("rag_vector", index)


def query_terms(text):
    return _tokenize(text)


def _tokenize(text):
    return [t for t in re.findall(r"[a-zA-Z0-9_./-]{2,}", str(text or "").lower()) if len(t) < 60]


def _hash_dim(term):
    digest = hashlib.blake2b(term.encode("utf-8"), digest_size=4).digest()
    return int.from_bytes(digest, "big") % VECTOR_DIMS


def _embed(text):
    vec = Counter()
    for term in _tokenize(text):
        vec[str(_hash_dim(term))] += 1
    norm = math.sqrt(sum(v * v for v in vec.values())) or 1.0
    return {k: round(v / norm, 6) for k, v in vec.items()}


def _cosine(left, right):
    if not left or not right:
        return 0.0
    if len(left) > len(right):
        left, right = right, left
    return sum(value * right.get(key, 0) for key, value in left.items())


def _import_status(module_name):
    try:
        __import__(module_name)
        return True
    except Exception:
        return False


def parser_status():
    return {
        "text": "enabled",
        "pdf": "enabled" if _import_status("pypdf") else "dependency_missing",
        "docx": "enabled" if _import_status("docx") else "dependency_missing",
        "xlsx": "enabled" if _import_status("openpyxl") else "dependency_missing",
    }


def _parser_status(path):
    ext = path.suffix.lower()
    if ext in TEXT_EXTS:
        return {"supported": True, "parser": "text", "reason": "", "citation_kind": "line"}
    if ext in DOCUMENT_EXTS:
        parser = ext.lstrip(".")
        status = parser_status().get(parser)
        return {
            "supported": status == "enabled",
            "parser": parser,
            "reason": "" if status == "enabled" else status,
            "citation_kind": "page" if parser == "pdf" else ("sheet_rows" if parser == "xlsx" else "line"),
        }
    return {"supported": False, "parser": "unsupported", "reason": "unsupported_extension"}


def _read_plain_text(path, status):
    if path.stat().st_size > MAX_TEXT_BYTES:
        return [], {"supported": False, "parser": "text", "reason": "too_large"}
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
            status = {**status, "reason": "encoding_replaced"}
        except Exception:
            return [], {"supported": False, "parser": "text", "reason": "decode_failed"}
    except Exception:
        return [], {"supported": False, "parser": "text", "reason": "read_failed"}
    return [{"text": text, "line_start": 1, "citation_kind": "line"}], status


def _read_pdf(path, status):
    if path.stat().st_size > MAX_DOCUMENT_BYTES:
        return [], {**status, "supported": False, "reason": "too_large"}
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        segments = []
        for index, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            text = text.strip()
            if text:
                page_no = index + 1
                segments.append({
                    "text": text,
                    "line_start": 1,
                    "page_start": page_no,
                    "page_end": page_no,
                    "citation_kind": "page",
                })
        if not segments:
            return [], {**status, "supported": False, "reason": "no_extractable_text"}
        return segments, status
    except Exception:
        return [], {**status, "supported": False, "reason": "parse_failed"}


def _read_docx(path, status):
    if path.stat().st_size > MAX_DOCUMENT_BYTES:
        return [], {**status, "supported": False, "reason": "too_large"}
    try:
        from docx import Document
        doc = Document(str(path))
        lines = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
                if values:
                    lines.append(" | ".join(values))
        text = "\n".join(lines).strip()
        if not text:
            return [], {**status, "supported": False, "reason": "no_extractable_text"}
        return [{"text": text, "line_start": 1, "citation_kind": "line"}], status
    except Exception:
        return [], {**status, "supported": False, "reason": "parse_failed"}


def _read_xlsx(path, status):
    if path.stat().st_size > MAX_DOCUMENT_BYTES:
        return [], {**status, "supported": False, "reason": "too_large"}
    try:
        from openpyxl import load_workbook
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        segments = []
        for sheet in workbook.worksheets:
            lines = []
            row_numbers = []
            for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                if row_index > MAX_XLSX_ROWS_PER_SHEET:
                    break
                values = [str(value).strip() for value in row if value is not None and str(value).strip()]
                if values:
                    lines.append(f"Row {row_index}: " + " | ".join(values))
                    row_numbers.append(row_index)
            if lines:
                segments.append({
                    "text": "\n".join(lines),
                    "line_start": row_numbers[0] if row_numbers else 1,
                    "line_numbers": row_numbers,
                    "sheet_name": sheet.title,
                    "citation_kind": "sheet_rows",
                })
        workbook.close()
        if not segments:
            return [], {**status, "supported": False, "reason": "no_extractable_text"}
        return segments, status
    except Exception:
        return [], {**status, "supported": False, "reason": "parse_failed"}


def _read_document(path):
    status = _parser_status(path)
    if not status["supported"]:
        return [], status
    parser = status.get("parser")
    if parser == "text":
        return _read_plain_text(path, status)
    if parser == "pdf":
        return _read_pdf(path, status)
    if parser == "docx":
        return _read_docx(path, status)
    if parser == "xlsx":
        return _read_xlsx(path, status)
    return [], {"supported": False, "parser": parser or "unsupported", "reason": "unsupported_extension"}


def _chunk_lines(text, lines_per_chunk=80, overlap=12, base_line=1, extra=None, line_numbers=None):
    lines = str(text or "").splitlines()
    if not lines:
        return []
    chunks = []
    step = max(1, lines_per_chunk - overlap)
    for start in range(0, len(lines), step):
        part = lines[start:start + lines_per_chunk]
        if not part:
            break
        joined = "\n".join(part).strip()
        if joined:
            if line_numbers and start < len(line_numbers):
                line_start = line_numbers[start]
                line_end = line_numbers[min(start + len(part) - 1, len(line_numbers) - 1)]
            else:
                line_start = base_line + start
                line_end = base_line + start + len(part) - 1
            chunk = {
                "text": joined,
                "line_start": line_start,
                "line_end": line_end,
                **(extra or {}),
            }
            if chunk.get("citation_kind") == "sheet_rows":
                chunk["row_start"] = line_start
                chunk["row_end"] = line_end
            chunks.append(chunk)
        if start + lines_per_chunk >= len(lines):
            break
    return chunks


def _chunk_segments(segments):
    chunks = []
    for segment in segments:
        extra = {
            key: segment.get(key)
            for key in ["page_start", "page_end", "sheet_name", "citation_kind"]
            if segment.get(key) not in {None, ""}
        }
        chunks.extend(_chunk_lines(
            segment.get("text", ""),
            base_line=segment.get("line_start") or 1,
            extra=extra,
            line_numbers=segment.get("line_numbers"),
        ))
    return chunks


def _walk(target):
    if target.is_file():
        return [target]
    files = []
    for p in target.rglob("*"):
        if any(x in SKIP_DIRS for x in p.parts):
            continue
        if p.is_file() and p.suffix.lower() not in SKIP_EXTS:
            if p.resolve() == INDEX_FILE.resolve():
                continue
            files.append(p)
    return files


def _scope(path="."):
    active = workspace.get_active()
    target = workspace.resolve_path(path or active.get("active_path") or ".")
    item = workspace.workspace_for_path(target)
    if not item:
        raise ValueError("path outside approved workspaces")
    root = Path(item.get("absolute_path") or workspace.resolve_path(item.get("id"))).resolve()
    if root not in target.parents and target != root:
        raise ValueError("path outside approved workspace")
    return target, item, root


def _source(file_path, item, root):
    rel = str(file_path.relative_to(root)).replace("\\", "/")
    wid = item.get("id") or "workspace"
    return f"{wid}/{rel}"


def ingest(path=".", label=None):
    target, item, root = _scope(path)
    index = _load()
    files = _walk(target)
    touched = []
    unchanged = []
    new_chunks = []
    docs = []
    skipped = []
    existing = {doc.get("source"): doc for doc in index.get("docs", [])}

    for file_path in files:
        source = _source(file_path, item, root)
        rel = str(file_path.relative_to(root)).replace("\\", "/")
        stat = file_path.stat()
        existing_doc = existing.get(source)
        if existing_doc and existing_doc.get("mtime") == stat.st_mtime and existing_doc.get("bytes") == stat.st_size:
            unchanged.append(source)
            continue
        segments, parser = _read_document(file_path)
        if not segments:
            skipped.append({"path": rel, "reason": parser.get("reason"), "parser": parser.get("parser")})
            continue
        touched.append(source)
        text = "\n\n".join(segment.get("text", "") for segment in segments)
        content_hash = hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()
        doc = {
            "source": source,
            "workspace_id": item.get("id"),
            "workspace_name": item.get("name"),
            "path": rel,
            "label": label or rel,
            "mtime": stat.st_mtime,
            "bytes": stat.st_size,
            "content_hash": content_hash,
            "parser": parser.get("parser"),
            "parser_status": parser.get("reason") or "ok",
            "citation_kind": parser.get("citation_kind") or "line",
        }
        docs.append(doc)
        for n, chunk in enumerate(_chunk_segments(segments)):
            chunk_text = chunk.get("text", "")[:5000]
            terms = Counter(_tokenize(chunk_text))
            if not terms:
                continue
            new_chunks.append({
                "id": f"{source}#{n}",
                "source": source,
                "workspace_id": item.get("id"),
                "path": rel,
                "chunk": n,
                "line_start": chunk.get("line_start"),
                "line_end": chunk.get("line_end"),
                "page_start": chunk.get("page_start"),
                "page_end": chunk.get("page_end"),
                "sheet_name": chunk.get("sheet_name"),
                "row_start": chunk.get("row_start"),
                "row_end": chunk.get("row_end"),
                "citation_kind": chunk.get("citation_kind") or parser.get("citation_kind") or "line",
                "parser": parser.get("parser"),
                "text": chunk_text,
                "terms": dict(terms),
                "term_count": sum(terms.values()),
                "vector": _embed(chunk_text),
                "mtime": stat.st_mtime,
                "content_hash": content_hash,
            })

    kept_docs = [d for d in index["docs"] if d.get("source") not in touched]
    kept_chunks = [c for c in index["chunks"] if c.get("source") not in touched]
    index["docs"] = kept_docs + docs
    index["chunks"] = kept_chunks + new_chunks
    _save(index)
    workspace.mark_indexed(item.get("id"), True)
    return {
        "path": str(target),
        "workspace_id": item.get("id"),
        "files_seen": len(files),
        "docs_indexed": len(docs),
        "docs_skipped_unchanged": len(unchanged),
        "docs_skipped": len(skipped),
        "skipped": skipped[:50],
        "chunks_indexed": len(new_chunks),
        "total_docs": len(index["docs"]),
        "total_chunks": len(index["chunks"]),
        "index_backend": "local-hash-vector-json",
        "parser_status": parser_status(),
    }


def stats():
    index = _load()
    return {
        "version": index.get("version", 2),
        "index_backend": "local-hash-vector-json",
        "docs": len(index["docs"]),
        "chunks": len(index["chunks"]),
        "updated_at": index.get("updated_at"),
        "sources": index["docs"][-25:],
        "parser_status": parser_status(),
    }


def _filter_chunks(chunks, path=None):
    if not path or path == ".":
        active = workspace.get_active()
        wid = active.get("active_id")
        return [c for c in chunks if c.get("workspace_id") == wid]
    target, item, root = _scope(path)
    prefix = str(target.relative_to(root)).replace("\\", "/") if target != root else ""
    out = []
    for chunk in chunks:
        if chunk.get("workspace_id") != item.get("id"):
            continue
        cpath = chunk.get("path", "")
        if not prefix or cpath == prefix or cpath.startswith(prefix + "/"):
            out.append(chunk)
    return out


def _path_score(query, q_terms, chunk):
    raw_query = str(query or "").strip().lower().replace("\\", "/")
    path = str(chunk.get("path") or "").lower().replace("\\", "/")
    source = str(chunk.get("source") or "").lower().replace("\\", "/")
    filename = Path(path).name.lower()
    if not raw_query:
        return 0.0
    score = 0.0
    if raw_query == filename:
        score += 12.0
    if raw_query == path or raw_query == source:
        score += 14.0
    if raw_query and raw_query in path:
        score += 6.0
    elif raw_query and raw_query in source:
        score += 4.0
    path_terms = Counter(_tokenize(f"{path} {source} {filename}"))
    for term, q_count in q_terms.items():
        if path_terms.get(term):
            score += min(3.0, q_count * path_terms[term] * 0.8)
    return score


def search(query, limit=6, path=None):
    index = _load()
    q_terms = Counter(_tokenize(query))
    q_vec = _embed(query)
    if not q_terms:
        return {"query": query, "hits": []}

    chunks = _filter_chunks(index["chunks"], path)
    doc_freq = Counter()
    for chunk in chunks:
        for term in chunk.get("terms", {}):
            doc_freq[term] += 1
    total = max(1, len(chunks))
    scored = []
    for chunk in chunks:
        terms = chunk.get("terms", {})
        lexical = 0.0
        for term, q_count in q_terms.items():
            c_count = terms.get(term, 0)
            if not c_count:
                continue
            idf = math.log((1 + total) / (1 + doc_freq[term])) + 1
            lexical += (q_count * c_count * idf) / max(1, chunk.get("term_count", 1))
        semantic = _cosine(q_vec, chunk.get("vector", {}))
        path_match = _path_score(query, q_terms, chunk)
        score = lexical + (semantic * 0.35) + path_match
        if score > 0:
            scored.append((score, chunk, lexical, semantic, path_match))
    scored.sort(key=lambda x: x[0], reverse=True)
    hits = []
    for score, chunk, lexical, semantic, path_match in scored[:max(1, min(int(limit), 20))]:
        hits.append({
            "score": round(score, 5),
            "lexical_score": round(lexical, 5),
            "semantic_score": round(semantic, 5),
            "path_score": round(path_match, 5),
            "source": chunk["source"],
            "workspace_id": chunk.get("workspace_id"),
            "path": chunk.get("path"),
            "chunk": chunk["chunk"],
            "line_start": chunk.get("line_start"),
            "line_end": chunk.get("line_end"),
            "page_start": chunk.get("page_start"),
            "page_end": chunk.get("page_end"),
            "sheet_name": chunk.get("sheet_name"),
            "row_start": chunk.get("row_start"),
            "row_end": chunk.get("row_end"),
            "citation_kind": chunk.get("citation_kind") or "line",
            "parser": chunk.get("parser"),
            "mtime": chunk.get("mtime"),
            "text": chunk["text"],
            "citation": {
                "workspace_id": chunk.get("workspace_id"),
                "path": chunk.get("path"),
                "chunk": chunk.get("chunk"),
                "line_start": chunk.get("line_start"),
                "line_end": chunk.get("line_end"),
                "page_start": chunk.get("page_start"),
                "page_end": chunk.get("page_end"),
                "sheet_name": chunk.get("sheet_name"),
                "row_start": chunk.get("row_start"),
                "row_end": chunk.get("row_end"),
                "citation_kind": chunk.get("citation_kind") or "line",
                "mtime": chunk.get("mtime"),
            },
        })
    return {"query": query, "hits": hits}


def chunks_for_path(path=None):
    index = _load()
    return _filter_chunks(index["chunks"], path)


def reset():
    _save(_blank())
    return stats()
